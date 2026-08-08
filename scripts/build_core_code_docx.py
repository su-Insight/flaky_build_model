from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "docs" / "核心源代码-软著提交版-紧凑版.docx"

SOFTWARE_NAME = "持续集成不稳定构建检测系统"
SOFTWARE_SHORT_NAME = "CIBD"
SOFTWARE_VERSION = "V1.0"


def set_run_font(run, *, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(f"{SOFTWARE_NAME} {SOFTWARE_VERSION} 核心源代码")
    set_run_font(header_run, name="Microsoft YaHei", size=9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_before = footer.add_run("第 ")
    set_run_font(run_before, name="Microsoft YaHei", size=9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    inner_run = OxmlElement("w:r")
    inner_text = OxmlElement("w:t")
    inner_text.text = "1"
    inner_run.append(inner_text)
    field.append(inner_run)
    footer._p.append(field)
    run_after = footer.add_run(" 页")
    set_run_font(run_after, name="Microsoft YaHei", size=9)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{SOFTWARE_NAME}[简称：{SOFTWARE_SHORT_NAME}] {SOFTWARE_VERSION}")
    set_run_font(r, name="Microsoft YaHei", size=18, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("核心源代码文档")
    set_run_font(r2, name="Microsoft YaHei", size=15, bold=True)

    rows = [
        ("软件名称", SOFTWARE_NAME),
        ("软件简称", SOFTWARE_SHORT_NAME),
        ("版本号", SOFTWARE_VERSION),
        ("代码范围", "run.py 与 src/ 目录下核心 Python 源文件"),
        ("说明", "本文件用于整理软件著作权申请所需的核心程序源码，不包含测试代码、日志文件、结果文件及临时脚本。"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, name="Microsoft YaHei", size=10)

    doc.add_page_break()


def core_source_files() -> list[Path]:
    return [
        ROOT / "run.py",
        *sorted((ROOT / "src").rglob("*.py")),
    ]


def add_file_block(doc: Document, file_path: Path) -> None:
    relative_path = file_path.relative_to(ROOT)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(1)
    heading_run = heading.add_run(f"文件：{relative_path}")
    set_run_font(heading_run, name="Microsoft YaHei", size=12, bold=True)

    lines = file_path.read_text(encoding="utf-8").splitlines()
    for line_no, line in enumerate(lines, start=1):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f"{line_no:04d}  {line}")
        set_run_font(r, name="Courier New", size=8.5)


def build_docx() -> Path:
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    files = core_source_files()
    for index, file_path in enumerate(files):
        add_file_block(doc, file_path)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    path = build_docx()
    print(path)


if __name__ == "__main__":
    main()
